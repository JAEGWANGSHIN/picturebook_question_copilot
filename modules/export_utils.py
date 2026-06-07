"""Export helpers for Markdown and DOCX."""

from __future__ import annotations

import io
import json
import re
from datetime import datetime
from typing import Any


def slugify(text: str, default: str = "picturebook") -> str:
    text = (text or default).strip()
    text = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", text)
    text = re.sub(r"_+", "_", text).strip("_")
    return text or default


def dated_filename(title: str, suffix: str, ext: str) -> str:
    date = datetime.now().strftime("%Y%m%d")
    return f"{date}_{slugify(title)}_{suffix}.{ext}"


def markdown_table(rows: list[dict[str, Any]]) -> str:
    if not rows:
        return ""
    headers = list(rows[0].keys())
    out = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in rows:
        vals = []
        for h in headers:
            v = row.get(h, "")
            if isinstance(v, (list, dict)):
                v = json.dumps(v, ensure_ascii=False)
            vals.append(str(v).replace("\n", "<br>").replace("|", "/"))
        out.append("| " + " | ".join(vals) + " |")
    return "\n".join(out)


def result_to_markdown(title: str, result_type: str, result: Any) -> str:
    title = title or "그림책 질문수업 결과물"
    md = [f"# {title}", "", f"- 결과물 유형: {result_type}", f"- 생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}", ""]

    if isinstance(result, list):
        if all(isinstance(x, dict) for x in result):
            md.append(markdown_table(result))
        else:
            md.extend([f"- {x}" for x in result])
    elif isinstance(result, dict):
        if "metadata" in result and "steps" in result:
            md.append("## 수업 개요")
            for k, v in result["metadata"].items():
                if isinstance(v, list):
                    md.append(f"- **{k}**: " + ", ".join(map(str, v)))
                else:
                    md.append(f"- **{k}**: {v}")
            md.append("\n## 수업 흐름")
            md.append(markdown_table(result["steps"]))
        else:
            for k, v in result.items():
                md.append(f"## {k}")
                if isinstance(v, list):
                    if v and all(isinstance(x, dict) for x in v):
                        md.append(markdown_table(v))
                    else:
                        md.extend([f"- {x}" for x in v])
                elif isinstance(v, dict):
                    md.append("```json")
                    md.append(json.dumps(v, ensure_ascii=False, indent=2))
                    md.append("```")
                else:
                    md.append(str(v))
                md.append("")
    else:
        md.append(str(result))

    md.append("\n---\n저작권 안내: 이 결과물은 원문 복제가 아니라 수업 설계용 요약·질문·활동지입니다.")
    md.append("정서 안전 안내: 학생의 감정을 진단하지 않고, 개인 경험 공개는 선택형으로 운영합니다.")
    return "\n".join(md).strip()


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    for line in markdown_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:].strip(), level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:].strip(), level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:].strip(), level=3)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        elif stripped.startswith("|"):
            # Keep Markdown tables as monospaced text for reliability in MVP.
            doc.add_paragraph(stripped)
        elif stripped.startswith("---"):
            doc.add_paragraph("—" * 20)
        else:
            doc.add_paragraph(stripped)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
