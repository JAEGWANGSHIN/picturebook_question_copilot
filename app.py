from __future__ import annotations
import json, os, subprocess, tempfile
from datetime import datetime
from io import BytesIO

import streamlit as st
from openai import OpenAI
from docx import Document
from docx.shared import Pt

st.set_page_config(
    page_title="AI 그림책 질문수업 설계기",
    page_icon="📚",
    layout="centered",
    initial_sidebar_state="collapsed",
)

# ═══════════════════════════════════════════════════════════════════
# DB
# ═══════════════════════════════════════════════════════════════════
PICTUREBOOK_DB = [
    {"id":"pb001","title":"말놀이 동시집","author":"최승호·방시혁","theme":["음운인식","어휘"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"동물 이름과 의성어로 구성된 말놀이 동시 모음","literacy_elements":["음운인식","어휘"],"reason":"한국어 음소·음절 인식, 운율 체험"},
    {"id":"pb002","title":"수수께끼야 놀자","author":"이상교","theme":["음운인식","어휘"],"grade":["유치원","초등 1학년"],"summary":"수수께끼 형식으로 의미와 소리를 연결","literacy_elements":["음운인식"],"reason":"소리 단서로 단어 맞추기; 파닉스 연결"},
    {"id":"pb003","title":"Brown Bear, Brown Bear","author":"Bill Martin Jr.","theme":["음운인식","어휘"],"grade":["유치원","초등 1학년"],"summary":"색깔과 동물 이름이 반복되는 패턴 그림책","literacy_elements":["음운인식","어휘"],"reason":"반복 패턴으로 운율·예측 읽기"},
    {"id":"pb004","title":"단어 수집가","author":"Peter H. Reynolds","theme":["어휘","정체성"],"grade":["초등 1학년","초등 2학년","초등 3학년"],"summary":"소년이 세상의 단어들을 수집하는 이야기","literacy_elements":["어휘","음운인식"],"reason":"단어 가치 인식; 어휘 수집 동기 부여"},
    {"id":"pb005","title":"알사탕","author":"백희나","theme":["어휘","감정 이해","가족"],"grade":["초등 1학년","초등 2학년"],"summary":"신비한 사탕을 먹으면 주변의 소리가 들린다","literacy_elements":["어휘","감정"],"reason":"감각어·감정 어휘 풍부; 서정적 표현"},
    {"id":"pb006","title":"구름빵","author":"백희나","theme":["어휘","가족","상상력"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"비 오는 날 구름으로 빵을 만들어 날아다니는 이야기","literacy_elements":["어휘","이야기이해"],"reason":"요리·자연 어휘; 판타지 어휘 확장"},
    {"id":"pb007","title":"수박 수영장","author":"안녕달","theme":["어휘","배경지식","다양성 존중","상상력"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"커다란 수박 속 수영장에서 동네 사람들이 함께 노는 상상","literacy_elements":["어휘","배경지식"],"reason":"여름·공동체·감각 어휘; 배경지식 확장"},
    {"id":"pb008","title":"이상한 손님","author":"백희나","theme":["어휘","추론하기","배려"],"grade":["초등 1학년","초등 2학년"],"summary":"비 오는 날 하늘 나라에서 길 잃은 아이가 찾아온 이야기","literacy_elements":["어휘","추론"],"reason":"감정 어휘·비유 표현; 상황 맥락 어휘 추론"},
    {"id":"pb009","title":"내 귀는 짝짝이","author":"율리 슈타르크","theme":["자존감","다양성 존중","정체성"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"귀가 다른 토끼가 자신을 있는 그대로 받아들이는 이야기","literacy_elements":["감정","자존감"],"reason":"신체 다양성 수용; 자기 긍정 감정 어휘"},
    {"id":"pb010","title":"너는 특별하단다","author":"맥스 루케이도","theme":["자존감","정체성"],"grade":["유치원","초등 1학년","초등 2학년","초등 3학년"],"summary":"작은 나무 사람 펀치넬로가 자신의 가치를 깨닫는 이야기","literacy_elements":["감정","이야기이해"],"reason":"자존감 언어; 평가와 자기 가치 탐색"},
    {"id":"pb011","title":"중요한 사실","author":"마가렛 와이즈 브라운","theme":["정체성","어휘"],"grade":["초등 1학년","초등 2학년"],"summary":"사물의 가장 중요한 특성에 대해 이야기하는 철학적 그림책","literacy_elements":["어휘","추론"],"reason":"핵심 특성 파악; 자아에 대한 질문 생성"},
    {"id":"pb012","title":"고슴도치 X","author":"에밀리 그레이벳","theme":["친구 관계","감정 이해","의사소통"],"grade":["초등 1학년","초등 2학년"],"summary":"편지를 쓰다가 계속 틀려서 X로 지워나가는 고슴도치 이야기","literacy_elements":["감정","이야기이해"],"reason":"감정 표현의 어려움; 쓰기와 감정 연결"},
    {"id":"pb013","title":"100만 번 산 고양이","author":"사노 요코","theme":["자존감","정체성","감정 이해"],"grade":["초등 2학년","초등 3학년","초등 4학년"],"summary":"100만 번을 살면서 진정한 사랑을 깨달은 고양이 이야기","literacy_elements":["이야기이해","감정"],"reason":"삶과 사랑의 의미; 감정 변화 추적"},
    {"id":"pb014","title":"빈집에 온 손님","author":"김유경","theme":["두려움","용기","상상력"],"grade":["초등 1학년","초등 2학년"],"summary":"홀로 집을 지키던 아이가 상상 속 손님을 맞이하는 이야기","literacy_elements":["추론","감정"],"reason":"두려움과 상상력; 감정 탐색"},
    {"id":"pb015","title":"무지개 물고기","author":"마르쿠스 피스터","theme":["어휘","배려","친구 관계"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"아름다운 비늘을 나눠주며 친구를 사귀는 물고기 이야기","literacy_elements":["어휘","감정"],"reason":"바닷속 어휘; 나눔·감정 표현 어휘"},
    {"id":"pb016","title":"강아지똥","author":"권정생","theme":["자존감","감정 이해","배려"],"grade":["초등 1학년","초등 2학년","초등 3학년"],"summary":"아무도 거들떠보지 않던 강아지똥이 민들레의 거름이 된다","literacy_elements":["어휘","감정"],"reason":"자연 어휘·존재 가치 어휘; 문학적 표현"},
    {"id":"pb017","title":"돼지책","author":"앤서니 브라운","theme":["가족","다양성 존중","의사소통"],"grade":["초등 2학년","초등 3학년","초등 4학년"],"summary":"혼자 집안일을 하던 엄마가 집을 나가고 가족이 돼지로 변한다","literacy_elements":["이야기이해","추론"],"reason":"인물 동기·감정 변화 추적; 시각 상징 분석"},
    {"id":"pb018","title":"꽃들에게 희망을","author":"트리나 폴러스","theme":["자존감","정체성","용기"],"grade":["초등 2학년","초등 3학년","초등 4학년"],"summary":"애벌레가 자아를 찾아 나비가 되는 우화","literacy_elements":["어휘","추론"],"reason":"삶의 의미 어휘; 변화·희망 주제 어휘"},
    {"id":"pb019","title":"고구마구마","author":"사이다","theme":["어휘","의사소통","감정 이해"],"grade":["유치원","초등 1학년"],"summary":"고구마가 '구마'라고만 말하는 반복 언어유희 그림책","literacy_elements":["어휘","음운인식"],"reason":"파닉스 연결; 반복 패턴 어휘 강화"},
    {"id":"pb020","title":"괜찮아","author":"최숙희","theme":["자존감","다양성 존중","감정 이해"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"서로 다른 모습도 괜찮다는 자기 수용 이야기","literacy_elements":["감정","어휘"],"reason":"자기 수용·다양성 감정 표현 어휘"},
    {"id":"pb021","title":"7년 동안의 잠","author":"박완서","theme":["배경지식","감정 이해"],"grade":["초등 2학년","초등 3학년"],"summary":"흉년 든 개미마을에 나타난 번데기를 둘러싼 이야기","literacy_elements":["이야기이해","배경지식"],"reason":"생태 배경지식; 발단·전개·결말 구조 분석"},
    {"id":"pb022","title":"나쁜 어린이 표","author":"황선미","theme":["자존감","감정 이해"],"grade":["초등 1학년","초등 2학년","초등 3학년"],"summary":"잘못을 저지른 어린이가 표를 붙이고 다니는 이야기","literacy_elements":["이야기이해","감정"],"reason":"원인-결과; 인물 내면 변화 이해"},
    {"id":"pb023","title":"선물","author":"이수지","theme":["상상력","감정 이해"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"눈이 오는 날의 감동을 글 없이 그림만으로 표현","literacy_elements":["추론"],"reason":"무자(wordless) 그림책; 그림 추론 최적"},
    {"id":"pb024","title":"지각대장 존","author":"John Burningham","theme":["상상력","의사소통"],"grade":["초등 1학년","초등 2학년"],"summary":"매일 지각하는 존의 기상천외한 이유","literacy_elements":["이야기이해","추론"],"reason":"사실과 상상 구별; 인물 관점 이해"},
    {"id":"pb025","title":"100층짜리 집","author":"이와이 도시오","theme":["배경지식","상상력"],"grade":["유치원","초등 1학년"],"summary":"주인공이 100층까지 올라가며 여러 동물을 만나는 이야기","literacy_elements":["이야기이해","배경지식"],"reason":"순서·수 개념; 동물 생태 배경지식"},
    {"id":"pb026","title":"으뜸 헤엄이(Swimmy)","author":"Leo Lionni","theme":["친구 관계","용기","배려"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"혼자 헤엄치는 물고기가 친구들과 힘을 합치는 이야기","literacy_elements":["이야기 재구성","감정"],"reason":"시각적 장면 순서로 재구성 용이"},
    {"id":"pb027","title":"The Very Hungry Caterpillar","author":"Eric Carle","theme":["배경지식","어휘"],"grade":["유치원","초등 1학년"],"summary":"배고픈 애벌레가 다양한 음식을 먹으며 성장하는 이야기","literacy_elements":["이야기 재구성","배경지식"],"reason":"요일·음식·변태 순서 재구성; 반복 구조"},
    {"id":"pb028","title":"Where the Wild Things Are","author":"Maurice Sendak","theme":["감정 조절","상상력","가족"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"맥스가 상상의 세계로 여행하고 집으로 돌아오는 이야기","literacy_elements":["이야기 재구성","감정"],"reason":"여행 구조(출발-모험-귀환) 재구성 전형"},
    {"id":"pb029","title":"In My Heart","author":"Jo Witek","theme":["감정 이해","감정 조절"],"grade":["유치원","초등 1학년"],"summary":"다양한 감정을 신체 감각으로 묘사하는 그림책","literacy_elements":["감정","어휘"],"reason":"감정 어휘 10가지 명시적 학습"},
    {"id":"pb030","title":"왜냐하면(Because)","author":"Mo Willems","theme":["상상력","배경지식"],"grade":["초등 1학년","초등 2학년"],"summary":"연쇄적 원인-결과로 이어지는 이야기","literacy_elements":["이야기이해"],"reason":"'왜?' 질문 구조를 시각적으로 보여줌"},
    {"id":"pb031","title":"Two Bad Ants","author":"Chris Van Allsburg","theme":["상상력","용기"],"grade":["초등 2학년","초등 3학년"],"summary":"두 개미가 설탕 그릇으로 모험을 떠나는 이야기","literacy_elements":["추론"],"reason":"개미 시점으로 보는 세상: 시각 추론의 정수"},
    {"id":"pb032","title":"Voices in the Park","author":"앤서니 브라운","theme":["다양성 존중","친구 관계","의사소통"],"grade":["초등 2학년","초등 3학년","초등 4학년"],"summary":"같은 공원 방문을 4명의 서로 다른 목소리로 이야기","literacy_elements":["추론","이야기이해"],"reason":"관점 추론; 같은 사건의 다른 해석"},
    {"id":"pb033","title":"선생님이 나를 모르면","author":"이상교","theme":["정체성","의사소통"],"grade":["초등 1학년"],"summary":"아이가 선생님에게 자신을 소개하는 이야기","literacy_elements":["이야기이해"],"reason":"나에 대한 질문 생성; 자기 이해 촉진"},
    {"id":"pb034","title":"나는 어떻게 생겨났을까?","author":"과학그림책","theme":["배경지식","정체성"],"grade":["초등 1학년","초등 2학년"],"summary":"탄생의 과학적 사실을 어린이 눈높이로 설명","literacy_elements":["배경지식"],"reason":"배경지식 궁금증에서 질문 생성 자연 유도"},
    {"id":"pb035","title":"The Invisible String","author":"Patrice Karst","theme":["감정 이해","두려움","가족"],"grade":["유치원","초등 1학년","초등 2학년"],"summary":"사랑하는 사람과의 보이지 않는 연결 이야기","literacy_elements":["감정"],"reason":"분리불안·연결감 감정; 저학년 적합"},
]

def db_search(theme="", grade=""):
    return [b for b in PICTUREBOOK_DB if
            ((not theme) or any(theme in t for t in b["theme"])) and
            ((not grade) or grade in b["grade"])]

def db_get_by_title(title):
    return next((b for b in PICTUREBOOK_DB if b["title"] == title), None)

def db_all_themes():
    t = set()
    for b in PICTUREBOOK_DB: t.update(b["theme"])
    return sorted(t)

# ═══════════════════════════════════════════════════════════════════
# CSS
# ═══════════════════════════════════════════════════════════════════
CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Gaegu:wght@700&family=Nanum+Gothic:wght@400;700;800&display=swap');

html, body,
[data-testid="stAppViewContainer"],
[data-testid="stAppViewContainer"] > .main { background-color: #FFF9F0 !important; }
[data-testid="stSidebar"] { display: none !important; }
.main .block-container { max-width: 740px !important; padding: 1.8rem 1.4rem 3rem !important; }

.app-header { text-align:center; padding:1.6rem 0 1.2rem; }
.app-icon   { font-size:2.4rem; display:block; margin-bottom:6px;
              animation:bob 3s ease-in-out infinite; }
@keyframes bob { 0%,100%{transform:translateY(0);} 50%{transform:translateY(-6px);} }
.app-title  { font-family:'Gaegu',cursive !important; font-size:clamp(1.6rem,5vw,2.2rem) !important;
              color:#3D2B1F !important; line-height:1.2 !important; margin:0 0 .3rem !important; }
.app-sub    { font-size:.88rem; color:#7D5A4A; font-weight:700 !important; }

.divider { height:2px;
  background:repeating-linear-gradient(90deg,#FFCC80 0,#FFCC80 8px,transparent 8px,transparent 14px);
  border:none; margin:1.4rem 0; }

.sec-label { font-family:'Gaegu',cursive !important; font-size:1.08rem !important;
             color:#5D3A1A !important; font-weight:700 !important;
             margin:0 0 .7rem !important; display:flex; align-items:center; gap:5px; }

/* 폼 */
[data-testid="stSelectbox"] > div > div,
[data-testid="stTextInput"] > div > div > input,
[data-testid="stTextArea"] > div > div > textarea {
  border-radius:10px !important; border:2px solid #E8C9A0 !important;
  background:#FFFDF7 !important; font-family:'Nanum Gothic',sans-serif !important; }
[data-testid="stSelectbox"] > div > div:focus-within,
[data-testid="stTextInput"] > div > div > input:focus,
[data-testid="stTextArea"] > div > div > textarea:focus {
  border-color:#FF7043 !important; box-shadow:0 0 0 3px #FF704320 !important; }

/* 책 카드 */
.book-card { background:linear-gradient(135deg,#FFF8E7,#FFF0F5);
             border:2px solid #FFCC80; border-radius:12px;
             padding:.85rem 1rem; margin:.4rem 0;
             display:flex; gap:10px; align-items:flex-start; }
.bc-icon  { font-size:1.5rem; flex-shrink:0; }
.bc-title { font-weight:800; color:#3D2B1F; font-size:.88rem; }
.bc-meta  { color:#9E8070; font-size:.76rem; margin-top:2px; line-height:1.4; }
.bc-tags  { margin-top:4px; display:flex; flex-wrap:wrap; gap:3px; }
.tag { background:#E8F5E9; color:#1B5E20; border:1px solid #A5D6A7;
       border-radius:20px; padding:1px 7px; font-size:.67rem; font-weight:800; }

/* AI 추천 칩 */
.rec-chips { display:flex; flex-wrap:wrap; gap:8px; margin-top:.8rem; }
.rec-chip {
  background:white; border:2px solid #E8C9A0; border-radius:50px;
  padding:6px 14px; font-size:.82rem; font-weight:700; color:#5D3A1A;
  cursor:pointer; transition:border-color .15s,background .15s;
  display:flex; align-items:center; gap:5px; }
.rec-chip:hover { border-color:#FF7043; background:#FFF3EE; }
.rec-chip .rn { font-weight:800; color:#FF7043; }

/* 질문 카드 그리드 */
.q-grid { display:grid; grid-template-columns:repeat(auto-fill,minmax(200px,1fr));
          gap:10px; margin-top:.6rem; }
.q-card { background:white; border:2px solid #F0D9B8; border-radius:12px;
          padding:.8rem .9rem; transition:transform .15s,box-shadow .15s; }
.q-card:hover { transform:translateY(-3px); box-shadow:0 6px 16px #00000014; }
.q-card .qt  { font-size:.67rem; font-weight:800; border-radius:20px;
               padding:2px 8px; display:inline-block; margin-bottom:5px; }
.q-card .qt.사실  { background:#E3F2FD; color:#1565C0; }
.q-card .qt.추론  { background:#F3E5F5; color:#6A1B9A; }
.q-card .qt.평가  { background:#E8F5E9; color:#1B5E20; }
.q-card .qt.감정  { background:#FCE4EC; color:#880E4F; }
.q-card .qt.작가  { background:#FFF3E0; color:#E65100; }
.q-card .qt.삶연결{ background:#F1F8E9; color:#33691E; }
.q-card .qtext { font-size:.82rem; color:#3D2B1F; line-height:1.5; }

/* 단계 버튼 */
.step-row { display:grid; grid-template-columns:repeat(4,1fr); gap:8px; margin:1rem 0; }
.step-btn-wrap button {
  border-radius:12px !important; font-family:'Gaegu',cursive !important;
  font-size:.85rem !important; font-weight:700 !important; }

/* 결과 아코디언 */
[data-testid="stExpander"] {
  background:white !important; border:2px solid #F0D9B8 !important;
  border-radius:12px !important; margin-bottom:6px !important; }
[data-testid="stExpander"]:hover { border-color:#FF7043 !important; }
[data-testid="stExpander"] summary {
  font-family:'Gaegu',cursive !important; font-size:.98rem !important;
  font-weight:700 !important; color:#3D2B1F !important; padding:.55rem .9rem !important; }

/* 메인 버튼 */
[data-testid="baseButton-primary"] {
  background:#FF7043 !important; border:none !important; border-radius:50px !important;
  font-family:'Gaegu',cursive !important; font-size:1.05rem !important;
  font-weight:700 !important; color:white !important;
  box-shadow:0 4px 0 #BF360C !important; letter-spacing:1px;
  transition:transform .1s,box-shadow .1s !important; }
[data-testid="baseButton-primary"]:hover  { transform:translateY(-2px) !important; box-shadow:0 6px 0 #BF360C !important; }
[data-testid="baseButton-primary"]:active { transform:translateY(2px) !important;  box-shadow:0 2px 0 #BF360C !important; }
[data-testid="baseButton-secondary"] {
  border-radius:50px !important; border:2px solid #E8C9A0 !important;
  background:white !important; font-weight:700 !important;
  transition:border-color .15s !important; }
[data-testid="baseButton-secondary"]:hover { border-color:#FF7043 !important; }

/* 탭 */
[data-testid="stTabs"] [role="tab"] {
  font-family:'Gaegu',cursive !important; font-size:.95rem !important;
  font-weight:700 !important; border-radius:10px 10px 0 0 !important; }
[data-testid="stTabs"] [role="tab"][aria-selected="true"] {
  color:#BF360C !important; border-bottom:3px solid #FF7043 !important; }

/* PPT 버튼 특별 색 */
.ppt-wrap [data-testid="baseButton-secondary"] {
  background:linear-gradient(135deg,#E8F5E9,#C8E6C9) !important;
  border:2px solid #81C784 !important; color:#1B5E20 !important; }

/* 반응형 */
@media(max-width:600px){
  .main .block-container{padding:1rem .8rem 2rem !important;}
  .app-title{font-size:1.4rem !important;}
  .q-grid{grid-template-columns:1fr;}
  .step-row{grid-template-columns:1fr 1fr;}
}
</style>
"""

PPTX_SCRIPT = os.path.join(os.path.dirname(__file__), "make_pptx.js")

# ═══════════════════════════════════════════════════════════════════
# OpenAI 헬퍼
# ═══════════════════════════════════════════════════════════════════
def get_client():
    key = st.secrets.get("OPENAI_API_KEY", None) or os.getenv("OPENAI_API_KEY")
    if not key:
        st.error("OpenAI API 키가 없습니다. Streamlit Secrets에 OPENAI_API_KEY를 등록하세요.")
        st.stop()
    return OpenAI(api_key=key)

def chat(system: str, user: str, max_tokens: int = 1200) -> str:
    client = get_client()
    r = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"system","content":system},
                  {"role":"user","content":user}],
        temperature=0.7,
        max_tokens=max_tokens,
    )
    return r.choices[0].message.content or ""

# ── AI 추천 ──────────────────────────────────────────────────────
def ai_recommend_books(situation: str) -> list[dict]:
    db_titles = "\n".join(
        f"- {b['title']} ({b['author']}): {b['summary']} [주제: {', '.join(b['theme'])}]"
        for b in PICTUREBOOK_DB
    )
    resp = chat(
        "당신은 초등 그림책 전문가입니다. 교사의 상황에 맞는 그림책을 추천하세요.",
        f"""교사 상황: {situation}

아래 그림책 목록 중 가장 적합한 책 4권을 추천하세요.
반드시 목록에 있는 제목만 사용하고, JSON 배열로만 응답하세요.
형식: [{{"title":"제목","reason":"한 문장 추천 이유"}}]

그림책 목록:
{db_titles}""",
        max_tokens=600,
    )
    try:
        import re
        m = re.search(r'\[.*\]', resp, re.DOTALL)
        return json.loads(m.group()) if m else []
    except Exception:
        return []

# ── 질문 생성 ─────────────────────────────────────────────────────
def gen_questions(grade, theme, book, book_info) -> dict:
    book_ctx = f"작가: {book_info['author']}\n줄거리: {book_info['summary']}" if book_info else ""
    resp = chat(
        "당신은 초등 그림책 질문 수업 전문가입니다. 반드시 JSON으로만 응답하세요.",
        f"""그림책: {book}
{book_ctx}
학년: {grade} / 주제: {theme}

읽기 전 질문 3개, 읽는 중 질문 5개, 읽은 후 질문 5개를 만들어 주세요.
각 질문은 사실/추론/평가/감정/작가/삶연결 중 하나의 유형을 가집니다.

반드시 아래 JSON 형식으로만 응답하세요:
{{"before":[{{"type":"유형","text":"질문"}},...],
  "during":[{{"type":"유형","text":"질문"}},...],
  "after" :[{{"type":"유형","text":"질문"}},...] }}""",
        max_tokens=1200,
    )
    try:
        import re
        m = re.search(r'\{.*\}', resp, re.DOTALL)
        return json.loads(m.group()) if m else {}
    except Exception:
        return {}

# ── 활동 생성 ─────────────────────────────────────────────────────
def gen_activities(grade, theme, book, lesson_time, student_context) -> str:
    return chat(
        "초등 수업 설계 전문가입니다. 한국어로 작성합니다.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book} / 시간:{lesson_time}
학생특성:{student_context or '없음'}

아래 형식으로 3가지 활동을 설계해 주세요.
## 활동 1: 도입 (활동명)
- 목표: ...
- 준비물: ...
- 진행: ...
- 교사 발문: ...

## 활동 2: 중심 (활동명)
[같은 형식]

## 활동 3: 정리 (활동명)
[같은 형식]""",
        max_tokens=1400,
    )

# ── 지도안 생성 ───────────────────────────────────────────────────
def gen_lessonplan(grade, theme, book, lesson_time, student_context) -> str:
    return chat(
        "초등 수업 설계 전문가입니다. 한국어로 작성합니다.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book} / 시간:{lesson_time}
학생특성:{student_context or '없음'}

도입/전개/정리 단계 지도안을 표 형식으로 작성하세요.
| 단계 | 시간 | 교사 활동 | 학생 활동 | 유의점 |
각 단계별로 2~3행씩 작성합니다.""",
        max_tokens=1000,
    )

# ── 평가+안내문 생성 ──────────────────────────────────────────────
def gen_eval_parent(grade, theme, book) -> str:
    return chat(
        "초등 수업 평가 및 학부모 소통 전문가입니다. 한국어로 작성합니다.",
        f"""학년:{grade} / 주제:{theme} / 그림책:{book}

아래 두 가지를 작성해 주세요.

## 평가 기준
관찰 평가 기준 4개 (충분/보통/노력 필요 3단계)
학생 자기평가 문항 3개

## 학부모 안내문
가정에 보내는 안내문 (그림책·주제 소개, 가정 대화 질문 3개, 따뜻한 어조)""",
        max_tokens=900,
    )

# ── DOCX ─────────────────────────────────────────────────────────
def make_docx(sections: dict, title: str) -> bytes:
    doc = Document()
    doc.styles["Normal"].font.name = "맑은 고딕"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"생성일: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    for sec_title, content in sections.items():
        if content:
            doc.add_heading(sec_title, level=2)
            for line in content.splitlines():
                s = line.strip()
                if not s: doc.add_paragraph("")
                elif s.startswith("## "): doc.add_heading(s[3:], level=3)
                elif s.startswith(("- ","• ")):
                    p = doc.add_paragraph(style="List Bullet"); p.add_run(s[2:])
                else: doc.add_paragraph(s)
    buf = BytesIO(); doc.save(buf); return buf.getvalue()

# ── PPTX ─────────────────────────────────────────────────────────
def make_pptx(grade, theme, book, lesson_time, questions, activities_text) -> bytes | None:
    # 활동 파싱
    acts = []
    if activities_text:
        import re
        for m in re.finditer(r'##\s+활동\s*\d+[:：]\s*(.+?)\n(.*?)(?=##\s+활동|\Z)',
                              activities_text, re.DOTALL):
            title_act = m.group(1).strip()
            body = m.group(2).strip()
            desc_m = re.search(r'진행[:：]\s*(.+?)(?:\n|$)', body)
            desc = desc_m.group(1).strip() if desc_m else body[:60]
            icons = ["🌱","📖","✍️"]
            acts.append({"icon": icons[len(acts) % 3], "title": title_act, "desc": desc})

    # 수업 목표 (간단 생성)
    obj_resp = chat(
        "한국어로 간결하게 답변하세요.",
        f"그림책 '{book}'로 {grade} {theme} 수업 목표 3개를 각각 한 문장으로 JSON 배열로만 답하세요. [\"목표1\",\"목표2\",\"목표3\"]",
        max_tokens=200,
    )
    try:
        import re
        m = re.search(r'\[.*?\]', obj_resp, re.DOTALL)
        objectives = json.loads(m.group()) if m else ["목표를 설정합니다.","질문을 만들어 봅니다.","생각을 표현합니다."]
    except Exception:
        objectives = ["목표를 설정합니다.","질문을 만들어 봅니다.","생각을 표현합니다."]

    data = {
        "title": f"「{book}」 질문수업",
        "subtitle": "AI 그림책 질문수업 설계안",
        "grade": grade,
        "theme": theme,
        "book": book,
        "lesson_time": lesson_time,
        "objectives": objectives,
        "questions": questions or {},
        "activities": acts or [
            {"icon":"🌱","title":"배경지식 활성화","desc":"그림책 표지 탐색, 경험 나누기"},
            {"icon":"📖","title":"대화형 읽기","desc":"PEER 절차로 그림책 읽기, 질문-응답"},
            {"icon":"✍️","title":"표현 활동","desc":"느낀 점 쓰기, 그림으로 표현하기"},
        ],
        "evaluations": [],
    }

    json_str = json.dumps(data, ensure_ascii=False)
    with tempfile.NamedTemporaryFile(suffix=".pptx", delete=False) as f:
        out_path = f.name

    try:
        result = subprocess.run(
            ["node", PPTX_SCRIPT, json_str, out_path],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode != 0:
            st.error(f"PPT 생성 오류: {result.stderr[:200]}")
            return None
        with open(out_path, "rb") as f:
            return f.read()
    except Exception as e:
        st.error(f"PPT 생성 실패: {e}")
        return None
    finally:
        try: os.unlink(out_path)
        except: pass

# ═══════════════════════════════════════════════════════════════════
# 질문 카드 렌더러
# ═══════════════════════════════════════════════════════════════════
def render_question_cards(questions: dict):
    sections = [
        ("before", "🌱 읽기 전"),
        ("during", "🔍 읽는 중"),
        ("after",  "💬 읽은 후"),
    ]
    for key, label in sections:
        qs = questions.get(key, [])
        if not qs: continue
        st.markdown(f"**{label}**")
        cols = st.columns(min(len(qs), 3))
        for i, q in enumerate(qs):
            qtype = q.get("type", "")
            qtext = q.get("text", q) if isinstance(q, dict) else q
            with cols[i % len(cols)]:
                st.markdown(
                    f'<div class="q-card">'
                    f'<span class="qt {qtype}">{qtype}</span>'
                    f'<div class="qtext">{qtext}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
        st.markdown("<br>", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════
def main():
    st.markdown(CSS, unsafe_allow_html=True)

    # ── 헤더 ──
    st.markdown("""
    <div class="app-header">
      <span class="app-icon">📚</span>
      <div class="app-title">AI 그림책 질문수업 설계기</div>
      <p class="app-sub">학년 · 주제 · 그림책을 고르면 수업 초안을 단계별로 만들어 드려요</p>
    </div>
    """, unsafe_allow_html=True)

    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)

    # ── STEP 1: 수업 조건 ──────────────────────────────────────────
    st.markdown('<div class="sec-label">① 수업 조건</div>', unsafe_allow_html=True)

    c1, c2, c3 = st.columns(3)
    with c1: grade = st.selectbox("학년", ["유치원","초등 1학년","초등 2학년","초등 3학년",
                                            "초등 4학년","초등 5학년","초등 6학년"],
                                   label_visibility="collapsed")
    with c2: theme = st.selectbox("주제", db_all_themes(), label_visibility="collapsed")
    with c3: lesson_time = st.selectbox("시간", ["40분","80분","120분","3차시","5차시"],
                                         label_visibility="collapsed")
    st.caption("📌 학년 　　 🎯 주제 　　 ⏰ 수업 시간")

    student_context = st.text_area("학생 특성 (선택)",
        placeholder="예: 1학년 입학 초기, 친구 관계가 서툰 편, 글쓰기 부담이 큼 등",
        height=64, label_visibility="collapsed")
    st.caption("📝 학생 특성 (선택)")

    # ── STEP 2: 그림책 선택 ────────────────────────────────────────
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-label">② 그림책 선택</div>', unsafe_allow_html=True)

    book_tab1, book_tab2, book_tab3 = st.tabs(["🤖 AI 추천", "📚 DB에서 찾기", "✏️ 직접 입력"])
    book = ""
    book_info = None

    # ─ AI 추천 탭 ─
    with book_tab1:
        situation_input = st.text_area(
            "우리 반 상황 입력",
            placeholder='예: "우리 반은 친구 관계 갈등이 많아요"\n"자존감이 낮아 자기 표현을 못하는 아이들이 있어요"',
            height=80, key="ai_situation", label_visibility="collapsed"
        )
        st.caption("📌 우리 반 상황을 자유롭게 입력하면 AI가 맞는 그림책을 추천해 드려요")

        if st.button("🤖 그림책 추천받기", key="btn_ai_rec"):
            if not situation_input.strip():
                st.warning("상황을 입력해 주세요.")
            else:
                with st.spinner("AI가 그림책을 고르는 중..."):
                    recs = ai_recommend_books(situation_input)
                st.session_state["ai_recs"] = recs

        if "ai_recs" in st.session_state and st.session_state["ai_recs"]:
            st.markdown("**추천 그림책** — 선택하면 바로 적용됩니다")
            chips_html = '<div class="rec-chips">'
            for i, r in enumerate(st.session_state["ai_recs"]):
                chips_html += (
                    f'<div class="rec-chip">'
                    f'<span class="rn">{i+1}</span> {r["title"]}'
                    f'</div>'
                )
            chips_html += "</div>"
            st.markdown(chips_html, unsafe_allow_html=True)

            chosen = st.radio("선택", [r["title"] for r in st.session_state["ai_recs"]],
                              label_visibility="collapsed", horizontal=True, key="ai_chosen")
            if chosen:
                book = chosen
                book_info = db_get_by_title(chosen)
                for r in st.session_state["ai_recs"]:
                    if r["title"] == chosen:
                        st.caption(f"💡 추천 이유: {r.get('reason','')}")
                        break

    # ─ DB 탭 ─
    with book_tab2:
        show_db = st.toggle("🔍 전체 DB 탐색 펼치기", value=False)
        if show_db:
            fc1, fc2 = st.columns(2)
            with fc1: ft = st.selectbox("주제", ["전체"] + db_all_themes(), key="dbt")
            with fc2: fg = st.selectbox("학년", ["전체","유치원","초등 1학년","초등 2학년",
                                                   "초등 3학년","초등 4학년"], key="dbg")
            filtered = db_search("" if ft=="전체" else ft, "" if fg=="전체" else fg)
            st.caption(f"검색 결과 {len(filtered)}권")
            cards = ""
            for b in filtered:
                tags = "".join(f'<span class="tag">{t}</span>' for t in b["theme"][:2])
                cards += (f'<div class="book-card" style="margin-bottom:6px;">'
                          f'<span class="bc-icon">📕</span>'
                          f'<div><div class="bc-title">{b["title"]}</div>'
                          f'<div class="bc-meta">{b["author"]}</div>'
                          f'<div class="bc-tags">{tags}</div></div></div>')
            st.markdown(cards, unsafe_allow_html=True)

        rec = db_search(theme=theme, grade=grade)
        if rec:
            sel = st.selectbox(f"추천 ({len(rec)}권 — {grade} × {theme})",
                               [b["title"] for b in rec])
            book = sel
            book_info = db_get_by_title(sel)
        else:
            st.info("조건에 맞는 책이 없습니다.")

    # ─ 직접 입력 탭 ─
    with book_tab3:
        custom = st.text_input("그림책 제목", placeholder="예: 알사탕", key="custom_book")
        if custom:
            book = custom
            book_info = db_get_by_title(custom)
            st.success("✅ DB에 있는 책입니다!" if book_info else "ℹ️ DB에 없는 책 — AI 일반 지식으로 진행합니다.")

    # 선택된 책 표시
    if book:
        if book_info:
            tags_html = "".join(f'<span class="tag">{e}</span>' for e in book_info["literacy_elements"])
            st.markdown(
                f'<div class="book-card">'
                f'<span class="bc-icon">📕</span>'
                f'<div><div class="bc-title">{book_info["title"]}</div>'
                f'<div class="bc-meta">{book_info["author"]} · {book_info["summary"]}</div>'
                f'<div class="bc-tags">{tags_html}</div></div></div>',
                unsafe_allow_html=True)
        else:
            st.info(f"📖 선택된 책: **{book}**")

    # ── STEP 3: 단계별 생성 ────────────────────────────────────────
    st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="sec-label">③ 단계별 생성</div>', unsafe_allow_html=True)

    if not book:
        st.info("책을 먼저 선택해 주세요.")
    else:
        sc1, sc2, sc3, sc4 = st.columns(4)
        with sc1:
            if st.button("❓ 질문 생성", use_container_width=True, key="btn_q"):
                with st.spinner("질문 생성 중..."):
                    qs = gen_questions(grade, theme, book, book_info)
                st.session_state["questions"] = qs
                st.session_state["q_meta"] = (grade, theme, book)

        with sc2:
            if st.button("🎨 활동 생성", use_container_width=True, key="btn_a"):
                with st.spinner("활동 생성 중..."):
                    acts = gen_activities(grade, theme, book, lesson_time, student_context)
                st.session_state["activities"] = acts

        with sc3:
            if st.button("🗒️ 지도안 생성", use_container_width=True, key="btn_l"):
                with st.spinner("지도안 생성 중..."):
                    lp = gen_lessonplan(grade, theme, book, lesson_time, student_context)
                st.session_state["lessonplan"] = lp

        with sc4:
            if st.button("⭐ 평가+안내문", use_container_width=True, key="btn_e"):
                with st.spinner("평가·안내문 생성 중..."):
                    ev = gen_eval_parent(grade, theme, book)
                st.session_state["eval_parent"] = ev

        # ── 결과 ──────────────────────────────────────────────────
        has_any = any(k in st.session_state for k in
                      ["questions","activities","lessonplan","eval_parent"])
        if has_any:
            st.markdown('<div class="divider"></div>', unsafe_allow_html=True)
            st.markdown('<div class="sec-label">④ 결과</div>', unsafe_allow_html=True)

            if "questions" in st.session_state and st.session_state["questions"]:
                with st.expander("❓ 질문 카드", expanded=True):
                    render_question_cards(st.session_state["questions"])

            if "activities" in st.session_state:
                with st.expander("🎨 활동 생성"):
                    st.markdown(st.session_state["activities"])

            if "lessonplan" in st.session_state:
                with st.expander("🗒️ 지도안"):
                    st.markdown(st.session_state["lessonplan"])

            if "eval_parent" in st.session_state:
                with st.expander("⭐ 평가 & 학부모 안내문"):
                    st.markdown(st.session_state["eval_parent"])

            # ── 다운로드 & PPT ───────────────────────────────────
            st.markdown("<br>", unsafe_allow_html=True)
            dl1, dl2, dl3 = st.columns([2, 2, 3])

            sections_for_docx = {
                "질문 생성": str(st.session_state.get("questions","")),
                "활동": st.session_state.get("activities",""),
                "지도안": st.session_state.get("lessonplan",""),
                "평가·학부모 안내문": st.session_state.get("eval_parent",""),
            }
            docx_title = f"{book}_{theme}_질문수업설계안"

            with dl1:
                st.download_button(
                    "📄 Word",
                    data=make_docx(sections_for_docx, docx_title),
                    file_name=f"{docx_title}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl2:
                combined_md = "\n\n".join(
                    f"## {k}\n{v}" for k, v in sections_for_docx.items() if v
                )
                st.download_button(
                    "📝 Markdown",
                    data=combined_md.encode("utf-8"),
                    file_name=f"{docx_title}.md",
                    mime="text/markdown",
                    use_container_width=True,
                )

            with dl3:
                st.markdown('<div class="ppt-wrap">', unsafe_allow_html=True)
                if st.button("🎞️ PPT 생성 (Canva·발표용)", use_container_width=True, key="btn_ppt"):
                    with st.spinner("🖍️ PPT 슬라이드 만드는 중..."):
                        pptx_bytes = make_pptx(
                            grade, theme, book, lesson_time,
                            st.session_state.get("questions", {}),
                            st.session_state.get("activities", ""),
                        )
                    if pptx_bytes:
                        st.session_state["pptx_bytes"] = pptx_bytes

                if "pptx_bytes" in st.session_state:
                    st.download_button(
                        "⬇️ PPT 다운로드",
                        data=st.session_state["pptx_bytes"],
                        file_name=f"{book}_{theme}_수업PPT.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                        use_container_width=True,
                    )
                st.markdown("</div>", unsafe_allow_html=True)


if __name__ == "__main__":
    main()
